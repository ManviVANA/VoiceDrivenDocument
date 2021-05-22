import subprocess
import wolframalpha
from docx2pdf import convert
import tkinter
import json
import random
import operator
import speech_recognition as sr
import pyttsx3 #text-to-speech
from docx import Document
import datetime
import wikipedia
import webbrowser
import os
import winshell
import pyjokes
import feedparser
import smtplib
import ctypes
import time
import requests
import shutil
from twilio.rest import Client
from ecapture import ecapture as ec
from bs4 import BeautifulSoup
from urllib.request import urlopen


engine = pyttsx3.init('sapi5')
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[0].id)

def speak(audio): #text to speech
    engine.say(audio)
    engine.runAndWait()

def takeCommand(): #speech to text
    r = sr.Recognizer()

    with sr.Microphone() as source:

        print("Listening...")

        #r.pause_threshold = 1
        audio = r.listen(source)

    try:
        print("Recognizing...")
        query = r.recognize_google(audio, language='en-in')
        print("User said:",query)

    except Exception as e:
        print(e)
        print("Unable to Recognize your voice.")
        return "None"

    return query



def Greetings():
    hour = int(datetime.datetime.now().hour)
    if hour >= 0 and hour < 12:
        speak("Good Morning!")
        print("Good Morning!")

    elif hour >= 12 and hour < 18:
        speak("Good Afternoon!")
        print("Good Afternoon!")

    else:
        speak("Good Evening!")
        print("Good Evening!")

    ass = ("Harry Potter")
    speak("Welcome to Voice Driven Document,I am here to Assist you,by the way I am ")
    print("Welcome to Voice Driven Document,I am here to Assist you,by the way I am ")
    speak(ass)
    print(ass)



def username():
    print("What should i call you then?")
    speak("What should i call you then?")
    uname = takeCommand()
    speak("Welcome")
    speak(uname)
    print("Welcome",uname,"!")
    print("Here is the list of Commands, I can work as your personal Assistant also")
    speak("Here is the list of Commands, I can work as your personal Assistant also")
    print("For editing a Document: open document")
    speak("For editing a Document: open document")
    print("For adding a table in Document: add table")
    speak("For adding a table in Document: add table")
    print("For adding information from Wikipedia: wikipedia")
    speak("For adding information from Wikipedia: wikipedia")
    print("For adding a note & displaying a note : "
          " add a note "
          " show a note ")
    speak("For adding a note & displaying a note : "
          " add a note "
          " show a note ")
    print("For adding Time: add time")
    speak("For adding Time: add time")
    print("want to add a snap: open camera")
    speak("want to add a snap: open camera")
    print("For calculating: calculate")
    speak("For calculating: calculate")
    print("For opening a document: place document")
    speak("For opening a document: place document")
    print("For opening a ppt: place ppt")
    speak("For opening a ppt: place ppt")
    print("Stuck in the middle of Work: "
          " open google "
          " open stackoverflow "
          " open youtube ")
    speak("Stuck in the middle of Work: "
          " open google "
          " open stackoverflow "
          " open youtube ")
    print("having heavy load: "
          " play music "
          " open youtube ")
    speak("having heavy load:"
          " play music "
          " open youtube ")
    print("For Exit: Exit")
    speak("For Exit: Exit")
    speak("How can i Help you?")
    print("how can i Help you?")



if __name__ == '__main__': #main function

    Greetings()
    username()

    while True:

        query = takeCommand().lower()

        # All the commands said by user will be
        # stored here in 'query' and will be
        # converted to lower case for easily
        # recognition of command
        print(query)
        if 'wikipedia' in query:
            speak('Searching Wikipedia...')
            print('Searching Wikipedia...')
            query = query.replace("wikipedia", "")
            results = wikipedia.summary(query, sentences=3)
            speak("According to Wikipedia")

            print("According to Wikipedia")
            print(results)
            speak(results)
            speak("Information added to document")
            print("Information added to document")
            document = Document()
            p = document.add_paragraph(results)
            document.save('demo1.docx')

        elif 'open youtube' in query:
            speak("Here you go to Youtube\n")
            print("Here you go to Youtube\n")
            webbrowser.open("youtube.com")

        elif 'open google' in query:
            speak("Here you go to Google\n")
            print("Here you go to Google\n")
            webbrowser.open("google.com")

        elif 'open stack overflow' in query:
            speak("Here you go to Stack Over flow.Happy coding")
            print("Here you go to Stack Over flow.Happy coding")
            webbrowser.open("stackoverflow.com")


        elif 'play music' in query or "play song" in query:
            speak("Here you go with music")
            print("Here you go with music")
            music_dir = "C:\\Users\\ManviVana\\PycharmProjects\\miniproject\\songs"
            songs = os.listdir(music_dir)
            random = os.startfile(os.path.join(music_dir, songs[1]))

        elif 'add time'in query:
            v=int(datetime.datetime.now().hour)

            speak(f"the time is {v}")
            print(f"the time is {v}")


        elif 'how are you' in query:
            speak("I am fine, Thank you")
            speak("How are you, Sir")

        elif 'fine' in query or "good" in query:
            speak("It's good to know that your fine")

        elif "change my name to" in query:
            query = query.replace("change my name to", "")
            ass = query

        elif "change name" in query:
            speak("What would you like to call me, Sir ")
            ass = takeCommand()
            speak("Thanks for naming me")

        elif "what's your name" in query or "What is your name" in query:
            speak("My friends call me")
            speak(ass)
            print("My friends call me", ass)

        elif 'exit' in query:
            speak("Thanks for giving me your time")
            exit()

        elif "who made you" in query or "who created you" in query:
            speak("I have been created by Madhavi.")

        elif 'joke' in query:
            speak(pyjokes.get_joke())
            print(pyjokes.get_joke())

        elif 'search' in query or 'play' in query:

            query = query.replace("search", "")
            query = query.replace("play", "")
            webbrowser.open(query)

        elif "who i am" in query:
            speak("If you talk then definately your human.")

        elif "why you came to world" in query:
            speak("Thanks to Madhu. further It's a secret")

        elif 'display ppt' in query:
            speak("opening Power Point presentation")
            power = r"C:\\Users\\ManviVana\\PycharmProjects\\miniproject\\Voice Assistant.pptx"
            os.startfile(power)

        elif 'is love' in query:
            speak("It is 7th sense that destroy all other senses")

        elif "who are you" in query:
            speak("I am your virtual assistant created by madhu")

        elif 'reason for you' in query:
            speak("I was created as a Miniproject by madhu ")


        elif 'lock window' in query:
            speak("locking the device")
            ctypes.windll.user32.LockWorkStation()

        elif 'shutdown system' in query:
            speak("Hold On a Sec ! Your system is on its way to shut down")
            subprocess.call('shutdown / p /f')



        elif "don't listen" in query or "stop listening" in query:
            speak("for how much time you want to stop jarvis from listening commands")
            a = int(takeCommand())
            time.sleep(a)
            print(a)

        elif "where is" in query:
            query = query.replace("where is", "")
            location = query
            speak("User asked to Locate")
            speak(location)
            webbrowser.open("https://www.google.nl / maps / place/" + location + "")

        elif "open camera" in query or "take a photo" in query:
            ec.capture(0, "Jarvis Camera ", "img.jpg")

        elif "restart" in query:
            subprocess.call(["shutdown", "/r"])

        elif "hibernate" in query or "sleep" in query:
            speak("Hibernating")
            subprocess.call("shutdown / h")

        elif "log off" in query or "sign out" in query:
            speak("Make sure all the application are closed before sign-out")
            time.sleep(5)
            subprocess.call(["shutdown", "/l"])

        elif "write a note" in query:
            speak("What should i write, sir")
            note = takeCommand()
            file = open('jarvis.txt', 'w')
            speak("Sir, Should i include date and time")
            snfm = takeCommand()
            if 'yes' in snfm or 'sure' in snfm:
                strTime = datetime.datetime.now().strftime("% H:% M:% S")
                file.write(strTime)
                file.write(" :- ")
                file.write(note)
            else:
                file.write(note)


        elif "show note" in query:
            speak("Showing Notes")
            file = open("jarvis.txt", "r")
            print(file.read())
            speak(file.read(6))


        elif "jarvis" in query:

            speak(ass)
            speak(" in your service")

        elif "open wikipedia" in query:
            webbrowser.open("wikipedia.com")

        elif "Good Morning" in query:
            speak("A warm" + query)
            speak("How are you Mister")
            speak(ass)

        # most asked question from google Assistant
        elif "will you be my gf" in query or "will you be my bf" in query:
            speak("I'm not sure about, may be you should give me some time")

        elif "how are you" in query:
            speak("I'm fine, glad you me that")

        elif "i love you" in query:
            speak("It's hard to understand")

        elif "what is" in query or "who is" in query:

            # Use the same API key
            # that we have generated earlier
            client = wolframalpha.Client("API_ID")
            res = client.query(query)

            try:
                print(next(res.results).text)
                speak(next(res.results).text)
            except StopIteration:
                print("No results")

        elif "Are you Married" in query:
            speak("Not yet")

        elif 'open document' in query:
            speak("Document is ready to edit")
            print("Document is ready to edit")

            document = Document()
            speak("what do you want edit in heading coloumn?")
            print("what do you want edit in heading coloumn?")
            n=takeCommand()
            t=document.add_heading(n,0)
            speak('heading added')
            print('heading added')
            print("\n")
            speak('what i have to add in paragraph?')
            print('what i have to add in paragraph?')
            x = takeCommand()
            p = document.add_paragraph()
            print("\n")
            speak("want to apply any styles"
                  "apply bold"
                  "apply italic")

            print("want to apply any styles ")
            print("apply bold ")
            print("apply italic")
            n2 = takeCommand()
            if n2 == 'apply bold':
                p.add_run(x).bold = True
                print("Styles added")
                speak('paragraph added')
                print('paragraph added')
            elif n2 == 'apply italic':
                p.add_run(x).italic = True
                print("Styles added")
                speak('paragraph added')
                print('paragraph added')
            else:
                document.add_paragraph(x)
                speak('paragraph added')
                print('paragraph added')
            document.save('demo1.docx')

        elif 'add table' in query:
            document = Document()
            row=3
            col=3
            table = document.add_table(row, col)

            hdr_cells = table.rows[0].cells
            speak("heading of each col")
            print("heading of each col")
            for i in range(col):
                b=takeCommand()
                hdr_cells[i].text = b
                print("added")
                speak("added")
            print("ok")
            speak("ok")
            speak("contents in columns one by one")
            print("contents in columns one by one")
            for i in range(3):
                for j in range(3):
                    b1 = takeCommand()
                    row_cells = table.add_row().cells
                    row_cells[i][j].text = b1
                    print("ok")
                    speak("added")
            print("ok")
            speak("ok")
            document.add_page_break()
            document.save('demo1.docx')
        elif 'place document' in query:
            speak("opening document")
            print("opening document")
            doc = r"C:\\Users\\ManviVana\\PycharmProjects\\miniproject\\demo1.docx"
            os.startfile(doc)
        elif 'convert into pdf' in query:
            speak("converting into pdf")
            print("converting into pdf")
            path=os.listdir(r"C:\\Users\\ManviVana\\PycharmProjects\\miniproject")
            doc1=r"C:\\Users\\ManviVana\\PycharmProjects\\miniproject\\demo1.docx"
            pdf1=r"C:\\Users\\ManviVana\\PycharmProjects\\miniproject\\demo1.pdf"

            convert(doc1,pdf1)
            speak("converted into pdf")
            print("converted into pdf")
        elif 'place pdf' in query:
            speak("opening pdf")
            print("opening pdf")
            doc = r"C:\\Users\\ManviVana\\PycharmProjects\\miniproject\\demo1.pdf"
            os.startfile(doc)
        elif 'print document' in query:
            speak('printing document')
            print('printing document')
            file='demo.pdf'
            os.startfile(file,'print')









    # elif "" in query:
    # Command go here
    # For adding more commands
